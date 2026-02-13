"""
Document Processing Pipeline for Knowledge Base.

Extracts text from documents, chunks it, generates embeddings, and stores in database.
"""

import asyncio
import logging
import uuid
from pathlib import Path

import pymupdf

from cachibot.models.knowledge import DocChunk, DocumentStatus
from cachibot.services.vector_store import VectorStore, get_vector_store
from cachibot.storage.repository import KnowledgeRepository

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Processes documents into searchable chunks with embeddings.

    Pipeline:
    1. Extract text from document (PDF, TXT, MD, DOCX)
    2. Split text into overlapping chunks
    3. Generate embeddings for each chunk
    4. Store chunks with embeddings in database
    """

    def __init__(
        self,
        chunk_size: int = 500,
        chunk_overlap: int = 50,
        vector_store: VectorStore | None = None,
    ):
        """
        Initialize the document processor.

        Args:
            chunk_size: Target number of words per chunk
            chunk_overlap: Number of overlapping words between chunks
            vector_store: VectorStore instance (uses singleton if not provided)
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self._vector_store = vector_store
        self._repo = KnowledgeRepository()

    async def _broadcast_status(
        self,
        bot_id: str,
        document_id: str,
        status: str,
        chunk_count: int | None = None,
    ) -> None:
        """Broadcast document status change via WebSocket."""
        try:
            from cachibot.api.websocket import get_ws_manager
            from cachibot.models.websocket import WSMessage

            ws = get_ws_manager()
            await ws.broadcast(
                WSMessage.document_status(
                    bot_id=bot_id,
                    document_id=document_id,
                    status=status,
                    chunk_count=chunk_count,
                )
            )
        except Exception as e:
            # Don't fail processing if broadcast fails
            logger.debug(f"Failed to broadcast document status: {e}")

    @property
    def vector_store(self) -> VectorStore:
        """Get the vector store (lazy initialization)."""
        if self._vector_store is None:
            self._vector_store = get_vector_store()
        return self._vector_store

    def extract_text(self, file_path: Path, file_type: str) -> str:
        """
        Extract text content from a document.

        Args:
            file_path: Path to the document file
            file_type: Type of file ('pdf', 'txt', 'md', 'docx')

        Returns:
            Extracted text content
        """
        if file_type == "pdf":
            return self._extract_pdf(file_path)
        elif file_type == "docx":
            return self._extract_docx(file_path)
        else:
            # txt, md - read as plain text
            return file_path.read_text(encoding="utf-8")

    def _extract_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF."""
        doc = pymupdf.open(file_path)
        try:
            # Join pages with form feed character (standard page separator)
            pages = []
            for page in doc:
                text = page.get_text()
                if text.strip():
                    pages.append(text)
            return "\n\n".join(pages)
        finally:
            doc.close()

    def _extract_docx(self, file_path: Path) -> str:
        """Extract text from DOCX using python-docx."""
        import docx

        doc = docx.Document(file_path)
        parts: list[str] = []

        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                parts.append(text)

        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n\n".join(parts)

    def chunk_text(self, text: str) -> list[str]:
        """
        Split text into overlapping chunks.

        Uses word-based chunking with configurable size and overlap.
        This preserves word boundaries and provides context overlap.

        Args:
            text: Full text to chunk

        Returns:
            List of text chunks
        """
        # Normalize whitespace
        text = " ".join(text.split())

        if not text:
            return []

        words = text.split()

        if len(words) <= self.chunk_size:
            return [text]

        chunks = []
        i = 0

        while i < len(words):
            # Get chunk_size words starting at position i
            chunk_words = words[i : i + self.chunk_size]
            chunk_text = " ".join(chunk_words)
            chunks.append(chunk_text)

            # Move forward by (chunk_size - overlap)
            i += self.chunk_size - self.chunk_overlap

            # If remaining words are less than overlap, stop
            if i >= len(words):
                break

        return chunks

    async def process_document(
        self,
        document_id: str,
        bot_id: str,
        file_path: Path,
        file_type: str,
    ) -> int:
        """
        Process a document through the full pipeline.

        Args:
            document_id: ID of the document record
            bot_id: Bot that owns this document
            file_path: Path to the uploaded file
            file_type: Type of file ('pdf', 'txt', 'md', 'docx')

        Returns:
            Number of chunks created

        Raises:
            Exception: If processing fails (document status will be set to FAILED)
        """
        try:
            logger.info(f"Processing document {document_id} for bot {bot_id}")

            # Broadcast processing start
            await self._broadcast_status(bot_id, document_id, "processing")

            # Step 1: Extract text
            logger.debug(f"Extracting text from {file_path}")
            text = await asyncio.get_event_loop().run_in_executor(
                None, self.extract_text, file_path, file_type
            )

            if not text.strip():
                logger.warning(f"No text extracted from document {document_id}")
                await self._repo.update_document_status(
                    document_id, DocumentStatus.READY, chunk_count=0
                )
                return 0

            # Step 2: Chunk text
            logger.debug("Chunking text")
            chunks_text = self.chunk_text(text)
            logger.info(f"Created {len(chunks_text)} chunks")

            if not chunks_text:
                await self._repo.update_document_status(
                    document_id, DocumentStatus.READY, chunk_count=0
                )
                return 0

            # Step 3: Generate embeddings
            logger.debug("Generating embeddings")
            embeddings = await self.vector_store.embed_texts(chunks_text)

            # Step 4: Create chunk records
            doc_chunks = [
                DocChunk(
                    id=str(uuid.uuid4()),
                    document_id=document_id,
                    bot_id=bot_id,
                    chunk_index=i,
                    content=chunk_text,
                    embedding=self.vector_store.serialize_embedding(embedding),
                )
                for i, (chunk_text, embedding) in enumerate(zip(chunks_text, embeddings))
            ]

            # Step 5: Store chunks
            logger.debug(f"Storing {len(doc_chunks)} chunks")
            await self._repo.save_chunks(doc_chunks)

            # Step 6: Update document status
            await self._repo.update_document_status(
                document_id, DocumentStatus.READY, chunk_count=len(doc_chunks)
            )

            logger.info(f"Document {document_id} processed: {len(doc_chunks)} chunks")

            # Broadcast ready status
            await self._broadcast_status(
                bot_id, document_id, "ready", chunk_count=len(doc_chunks)
            )

            return len(doc_chunks)

        except Exception as e:
            logger.error(f"Failed to process document {document_id}: {e}")
            await self._repo.update_document_status(document_id, DocumentStatus.FAILED)

            # Broadcast failure
            await self._broadcast_status(bot_id, document_id, "failed")

            raise


# Singleton instance
_processor: DocumentProcessor | None = None


def get_document_processor() -> DocumentProcessor:
    """Get the shared DocumentProcessor instance (config-aware)."""
    global _processor
    if _processor is None:
        from cachibot.config import Config

        config = Config.load()
        _processor = DocumentProcessor(
            chunk_size=config.knowledge.chunk_size,
            chunk_overlap=config.knowledge.chunk_overlap,
        )
    return _processor
